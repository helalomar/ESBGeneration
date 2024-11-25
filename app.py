#region---------------------------------------imports--------------------------------------------

import xml.etree.ElementTree    as ET
import xmlschema
import xml.dom.minidom          as md
import xml.dom.minidom
import os
import json
import xmlschema
import re
import sys
import xmltodict
from ast            import operator
from docx           import Document
from docx.shared    import Inches
from docx.oxml      import parse_xml
from datetime       import datetime
from docx.shared    import Inches
from docx.oxml.ns   import nsdecls
from itertools import islice, zip_longest
from flask import Flask, render_template, request, send_file
import zipfile
#endregion---------------------------imports-----------------------------------

#region-------------------------------------inputproccesing--------------------------------------


processed_elements = set()
iprocessed_elements = set()
rsprocessed_elements = set()
irsprocessed_elements = set()
visited_paths = set()
previousitem = 'globalizeditempreviously'
formatted_date = datetime.now().strftime('%Y-%m-%d')
xsd_attributes = {
    "xmlns:xsd": "http://www.w3.org/2001/XMLSchema",
    "xmlns:alr": "urn:example:yournamespace",
    "elementFormDefault": "qualified"
}
CreatorName = 'Omar Helal'
OperationMethod = 'Inquiry'
IP = '10.11.34.16'
PORT =  '443'

#region place holder declarations
svcrqplaceholder_text = "{RqBodyDec}"
svcrsplaceholder_text = "{RsBodyDec}"
svcrsbeplaceholder_text = "{RsBodyBE}"
Rqplaceholder_text = "{RqTableRow}"
placeholder_text = "{TableRow}"
#endregion

#region Create output folder and subfolders if they don't exist

xsd_folder = "XSD"
wsdl_folder = "WSDL"
sci_folder = "SCI"
sample_folder = "SAMPLE"
svc_folder = 'Service'
core_folder = 'Core'
BEName = 'BackendAPIName'

#endregion-------------------------------------------------------

#endregion---------------------------inputproccesing----------------------------

#region-------------------------------------functions--------------------------------------------
def getuserinput():
    global SvcID
    global SubSvcID
    global CategoryName
    global OperationName
    global functionality
    global message_entry
    global BECall
    global BEName
    global BE
    global app_folder
    global output_folder
    functionality = int(request.form.get("functionality", 1))
    BECall = request.form.get("BECall", '2')
    OperationName = request.form.get("OperationName", 'DefaultOp')
    output_folder = os.path.join(os.getcwd(), OperationName)
    CategoryName = request.form.get("CategoryName", 'AlRajhiDefault')
    SvcID = request.form.get("SvcID", '1000')
    SubSvcID = request.form.get("SubSvcID", '1000')
    message_entry = request.form['messageEntry']
    app_folder = f"{OperationName}_AppPrj"
    if functionality >= 4:
        BEName = request.form.get("BEName", 'BackendAPIName')
        BE = request.form.get("BE", 'BE')

def doworddoc():
    # Load the altered  Word document

    doc = Document(os.path.join(output_folder, sci_folder, f"{SvcID}-{SubSvcID}-SCI {OperationName}.docx"))
    # doc = Document("NeededFiles/TemplateDocx.docx")
    # Replace placeholders in the document
    for paragraph in doc.paragraphs:
        paragraph.text = paragraph.text.format(SvcID=SvcID, SubSvcID=SubSvcID, OperationName=OperationName,
                                            CategoryName=CategoryName, IP=IP, PORT=PORT,
                                            CreatorName=CreatorName, OperationMethod=OperationMethod,
                                            CreationDate=formatted_date)
    # Iterate through tables and replace place holders
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = cell.text.format(SvcID=SvcID, SubSvcID=SubSvcID, OperationName=OperationName,
                                            CategoryName=CategoryName, IP=IP, PORT=PORT,
                                            CreatorName=CreatorName, OperationMethod=OperationMethod,
                                            CreationDate=formatted_date, ReqBody=printsamplereq, ResBody=printsampleres,
                                            TableRow="", RqTableRow="")
    # Iterate through tables and adjust column widths
    for table in doc.tables:
        for column in table.columns:
            # Set the desired width (in inches) for each column
            column.width = Inches(1.0)  # Adjust this value according to your needs

    # Save the modified Word document to the output folder
    doc.save(f"{output_folder}/{sci_folder}/{SvcID}-{SubSvcID}-SCI {OperationName}.docx")

def dosvcandcoreapp():
    with open("NeededFiles/App/ServiceLib.esql", "r") as svclib_template_file:
        svclib_template = svclib_template_file.read().format(SvcID=SvcID, SubSvcID=SubSvcID, OperationName=OperationName,
                                                        CategoryName=CategoryName, IP=IP, PORT=PORT,
                                                        CreatorName=CreatorName, OperationMethod=OperationMethod,
                                                        BEName=BEName , BE=BE, CreationDate=formatted_date)

    with open(f"{output_folder}/{app_folder}/{svc_folder}/ServiceLib.esql", "w") as svclib_out_file:
        svclib_out_file.write(svclib_template)

    with open("NeededFiles/App/template.project", "r") as projectfile_template_file:
        projectfile_template = projectfile_template_file.read().format(OperationName=OperationName,
                                                        CategoryName=CategoryName)

    with open(f"{output_folder}/{app_folder}/a.project", "w") as projectfile_out_file:
        projectfile_out_file.write(projectfile_template)

    with open(svcesqltemp_path, "r") as svcesql_template_file:
        svcesql_template = svcesql_template_file.read().format(SvcID=SvcID, SubSvcID=SubSvcID, OperationName=OperationName,
                                                        CategoryName=CategoryName, IP=IP, PORT=PORT,
                                                        CreatorName=CreatorName, OperationMethod=OperationMethod,
                                                        BEName=BEName , BE=BE, CreationDate=formatted_date, RqBodyDec="", RsBodyDec="",RsBodyBE='ReplaceHere')

    with open(f"{output_folder}/{app_folder}/{svc_folder}/{OperationName}_Svc.esql", "w") as svcesql_out_file:
        svcesql_out_file.write(svcesql_template)

    with open("NeededFiles/App/Template_Svc.msgflow", "r") as svcmsgflow_template_file:
        svcmsgflow_template = svcmsgflow_template_file.read().format(SvcID=SvcID, SubSvcID=SubSvcID, OperationName=OperationName,
                                                        CategoryName=CategoryName, IP=IP, PORT=PORT,
                                                        CreatorName=CreatorName, OperationMethod=OperationMethod,
                                                        BEName=BEName , BE=BE)

    with open(f"{output_folder}/{app_folder}/{svc_folder}/{OperationName}_Svc.msgflow", "w") as svcmsgflow_out_file:
        svcmsgflow_out_file.write(svcmsgflow_template)

    with open("NeededFiles/App/Template.subflow", "r") as coresubflow_template_file:
        coresubflow_template = coresubflow_template_file.read().format(SvcID=SvcID, SubSvcID=SubSvcID, OperationName=OperationName,
                                                        CategoryName=CategoryName, IP=IP, PORT=PORT,
                                                        CreatorName=CreatorName, OperationMethod=OperationMethod,
                                                        BEName=BEName , BE=BE)

    with open(f"{output_folder}/{app_folder}/{core_folder}/{OperationName}.subflow", "w") as coresubflow_out_file:
        coresubflow_out_file.write(coresubflow_template)

    with open("NeededFiles/App/Template.esql", "r") as coreesql_template_file:
        coreesql_template = coreesql_template_file.read().format(SvcID=SvcID, SubSvcID=SubSvcID, OperationName=OperationName,
                                                        CategoryName=CategoryName, IP=IP, PORT=PORT,
                                                        CreatorName=CreatorName, OperationMethod=OperationMethod,
                                                        BEName=BEName , BE=BE, CreationDate=formatted_date)

    with open(f"{output_folder}/{app_folder}/{core_folder}/{OperationName}.esql", "w") as coreesql_out_file:
        coreesql_out_file.write(coreesql_template)

def dowsdlandxsd():

    #region function 2: XSD and WSDL Generation
    #request 
    xsd_root  = ET.Element("xsd:schema", xsd_attributes)
    xml_element = xml_t_rq
    generate_xsd_element(xsd_root, xml_element)
    xml_element = ET.ElementTree(xml_element)
    xsd_string = ET.tostring(xsd_root, encoding="utf-8").decode()
    formatted_xsd = md.parseString(xsd_string).toprettyxml(indent="    ")
    indexx = formatted_xsd.find('Body_Type"/>')
    modified_schema = formatted_xsd[indexx + 40:][:-(len("</xsd:schema>") + 1)]
    #response
    xsd_roott = ET.Element("xsd:schema", xsd_attributes)
    xml_elementr = xml_t_rs
    generate_xsd_element(xsd_roott, xml_elementr)
    xml_elementr = ET.ElementTree(xml_elementr)
    # Convert XSD element to string and format it
    xsd_stringg = ET.tostring(xsd_roott, encoding="utf-8").decode()
    formattedd_xsd = md.parseString(xsd_stringg).toprettyxml(indent="    ")
    indexx = formatted_xsd.find('Body_Type"/>')
    modifiedd_schema = formattedd_xsd[indexx + 40:][:-(len("</xsd:schema>") + 1)]
    #endregion 

    # Read and replace placeholders in template files
    with open("NeededFiles/Template.wsdl", "r") as wsdl_template_file:
        wsdl_template = wsdl_template_file.read().format( OperationName=OperationName,
                                                        CategoryName=CategoryName, IP=IP, PORT=PORT)

    with open("NeededFiles/Template.xsd", "r") as xsd_template_file:
        xsd_template = xsd_template_file.read().format( OperationName=OperationName,
                                                    CategoryName=CategoryName, IP=IP, PORT=PORT,
                                                    BodyHere=modified_schema, RBodyHere=modifiedd_schema)

    with open(f"{output_folder}/{wsdl_folder}/{CategoryName}.wsdl", "w") as wsdl_out_file:
        wsdl_out_file.write(wsdl_template)

    with open(f"{output_folder}/{xsd_folder}/{CategoryName}.xsd", "w") as xsd_out_file:
        xsd_out_file.write(xsd_template)

    if functionality >= 4:
        with open(f"{output_folder}/{app_folder}/{wsdl_folder}/{CategoryName}.wsdl", "w") as wsdl_out_file:
            wsdl_out_file.write(wsdl_template)

        with open(f"{output_folder}/{app_folder}/{xsd_folder}/{CategoryName}.xsd", "w") as xsd_out_file:
            xsd_out_file.write(xsd_template)

    with open("NeededFiles/CommonLibV2.xsd", "r") as clib_template:
        clib_out = clib_template.read()

    with open(f"{output_folder}/{xsd_folder}/CommonLibV2.xsd", "w") as Comxsd_out_file:
        Comxsd_out_file.write(clib_out)

def list_xml_fields_with_path_svcr(element, path=None):
    if path is None:
        path = []

    path.append(element.tag)  # Add current element to the path
    yield '.'.join(path), element.text  # Yield the current path and element text (if any)

    for child in element:
        yield from list_xml_fields_with_path_svcr(child, path.copy())  # Recursively traverse child elements

    if path:  # Check if the path list is not empty before backtracking
        path.pop()  # Remove the last element to backtrack in the path

def list_xml_fields_with_path_svc(element, path=None):
    if path is None:
        path = []

    path.append(element.tag)  # Add current element to the path
    yield '.*:'.join(path), element.text  # Yield the current path and element text (if any)

    for child in element:
        yield from list_xml_fields_with_path_svc(child, path.copy())  # Recursively traverse child elements

    if path:  # Check if the path list is not empty before backtracking
        path.pop()  # Remove the last element to backtrack in the path

def list_xml_fields_with_path(element, path=None):
    if path is None:
        path = []

    path.append(element.tag)  # Add current element to the path
    yield '/'.join(path), element.text  # Yield the current path and element text (if any)

    for child in element:
        yield from list_xml_fields_with_path(child, path.copy())  # Recursively traverse child elements

    if path:  # Check if the path list is not empty before backtracking
        path.pop()  # Remove the last element to backtrack in the path

def list_json_fields_with_path(element, path=None):
    if path is None:
        path = []

    current_path = '.'.join(path)
    if current_path not in visited_paths:
        if isinstance(element, dict):
            temp = 0
        elif isinstance(element, list):
            temp = 0
        else:
            yield current_path  # Yield the current path
            visited_paths.add(current_path)  # Add the current path to visited_paths set


    if isinstance(element, dict):
        for key, value in element.items():
            yield from list_json_fields_with_path(value, path + [key])  # Recursively traverse child elements
    elif isinstance(element, list):
        for index, item in enumerate(element):
            yield from list_json_fields_with_path(item, path + ['jsonlistitem'])  # Recursively traverse list items

def format_xml(xml_string):
    # Parse the XML string into an ElementTree object
    root = ET.fromstring(xml_string)
    
    # Function to indent XML elements recursively
    def indent(elem, level=0):
        i = "\n" + level * "  "
        if len(elem):
            if not elem.text or not elem.text.strip():
                elem.text = i + "  "
            if not elem.tail or not elem.tail.strip():
                elem.tail = i
            for subelem in elem:
                indent(subelem, level + 1)
            if not elem.tail or not elem.tail.strip():
                elem.tail = i
        else:
            if level and (not elem.tail or not elem.tail.strip()):
                elem.tail = i
    
    # Indent the XML elements
    indent(root)
    
    # Serialize the ElementTree back to a string
    formatted_xml = ET.tostring(root, encoding="unicode")
    
    return formatted_xml

def capitalize_xml_tags(element):
    # Capitalize first letter of each word in the tag name
    if BECall == '2':
        element.tag = re.sub(r"(?!^)([A-Z]+)", r"_\1", element.tag).title().replace("_", "")
    else:
        element.tag = element.tag[0].upper() + element.tag[1:]
    # Recursively capitalize XML tags for child elements
    for child in element:
        capitalize_xml_tags(child)

def remove_invalid_suffix(tree):
    for parent in tree.iter():
        children = list(parent)
        for child in children:
            # Check if child tag ends with 'LstItem'
            if child.tag.endswith('LstItem'):
                break
        else:
            # If none of the children end with 'LstItem', remove 'Lst' suffix from parent tag
            if parent.tag.endswith('Lst') or parent.tag.endswith('lst'):
                parent.tag = parent.tag[:-3]

def remove_invalid_suffix_t(element):
        for child in element:
            # Check if the element name contains the find_string
            if 'LstLstItem' in child.tag:
                # Modify the element name by replacing find_string with replace_string
                new_tag = child.tag.replace('LstLstItem', 'LstItem')
                child.tag = new_tag
            if 'LstLst' in child.tag:
                # Modify the element name by replacing find_string with replace_string
                new_tag = child.tag.replace('LstLst', 'Lst')
                child.tag = new_tag
            if 'ListLst' in child.tag:
                # Modify the element name by replacing find_string with replace_string
                new_tag = child.tag.replace('ListLst', 'Lst')
                child.tag = new_tag
            if 'ListsLst' in child.tag:
                # Modify the element name by replacing find_string with replace_string
                new_tag = child.tag.replace('ListsLst', 'sLst')
                child.tag = new_tag
            if 'Id' in child.tag:
                # Modify the element name by replacing find_string with replace_string
                new_tag = child.tag.replace('Id', 'ID')
                child.tag = new_tag
            if 'id' in child.tag:
                # Modify the element name by replacing find_string with replace_string
                new_tag = child.tag.replace('id', 'ID')
                child.tag = new_tag
            if 'Date' in child.tag:
                # Modify the element name by replacing find_string with replace_string
                new_tag = child.tag.replace('Date', 'Dt')
                child.tag = new_tag
            if 'Number' in child.tag:
                # Modify the element name by replacing find_string with replace_string
                new_tag = child.tag.replace('Number', 'Num')
                child.tag = new_tag
            if 'number' in child.tag:
                # Modify the element name by replacing find_string with replace_string
                new_tag = child.tag.replace('number', 'Num')
                child.tag = new_tag
            if 'Customer' in child.tag:
                # Modify the element name by replacing find_string with replace_string
                new_tag = child.tag.replace('Customer', 'Cust')
                child.tag = new_tag
            if 'Account' in child.tag:
                # Modify the element name by replacing find_string with replace_string
                new_tag = child.tag.replace('Account', 'Acct')
                child.tag = new_tag
            if 'Amount' in child.tag:
                # Modify the element name by replacing find_string with replace_string
                new_tag = child.tag.replace('Amount', 'Amt')
                child.tag = new_tag
            # Recursively modify child elements
            remove_invalid_suffix_t(child)

def zip_folder(folder_path, output_path):
    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for root, dirs, files in os.walk(folder_path):
            for file in files:
                file_path = os.path.join(root, file)
                zipf.write(file_path, os.path.relpath(file_path, folder_path))

def add_lines_after_placeholder(input_file_path, output_file_path, placeholder, lines_to_add, msgtype, value):
    global previousitem     
    with open(input_file_path, "r") as input_file:
        file_content = input_file.read()

    if placeholder in file_content:
        parts = file_content.split(placeholder)
        if  msgtype == 'rs':
            if value == None and ("Lst" not in lines_to_add):
                return
            elif str(lines_to_add)[-3:] == 'Lst' :
                modified_content = parts[0] + "\n\t\t"  + f"DECLARE J INTEGER CARDINALITY({BEName}Res." + svcrsbeplaceholder_text  + "[]);\n\t\t" + "DECLARE I INTEGER 1;\n\n\t\t" + "CREATE FIELD OutBody."+ lines_to_add.rsplit('.', 1)[-1]  + ";\n\t\t" + "DECLARE outLst REFERENCE TO OutBody."+ lines_to_add.rsplit('.')[-1]+ ";\n\n\t\t"+"WHILE I <= J DO\n\t\t\t" + placeholder + parts[1]  
            elif str(lines_to_add)[-4:] == 'Item' :
                if lines_to_add.rsplit('.', 1)[-1] in rsprocessed_elements:
                    return
                # Mark the element as processed
                rsprocessed_elements.add(lines_to_add.rsplit('.', 1)[-1])
                modified_content = parts[0] + "\n\t\t\t" f"DECLARE inLst REFERENCE TO {BEName}Res."+ svcrsbeplaceholder_text + "[I]\n\t\t\t" + "CREATE LASTCHILD OF outLst NAME '" +lines_to_add.rsplit('.', 1)[-1] + "';\n\t\t\t"  + "DECLARE outLstItem REFERENCE TO '" + lines_to_add.rsplit('.', 1)[-1]+ "[<];\n\t\t\t" + placeholder + parts[1]  
            elif "LstItem" in lines_to_add:
                if lines_to_add.rsplit('.', 1)[-1] in irsprocessed_elements:
                    return
                # Mark the element as processed
                irsprocessed_elements.add(lines_to_add.rsplit('.', 1)[-1])
                previousitem = lines_to_add
                modified_content = parts[0] + "SET outLstItem." + lines_to_add.rsplit('.', 1)[-1]  + " = inLst." + svcrsbeplaceholder_text + "\n\t\t\t"  + placeholder + parts[1] 
            elif "LstItem" in previousitem:
                modified_content = parts[0] + "SET I = I+1;\n\t\tEND WHILE;\n\n\t\t" + f"SET OutBody." + lines_to_add.split('.', 1)[-1]  + " = {BEName}Res." + svcrsbeplaceholder_text + "\n\t\t" + placeholder + parts[1]
                previousitem = 'globalizeditempreviously'
            else:
                modified_content = parts[0] + f"SET OutBody." + lines_to_add.split('.', 1)[-1]  + " = {BEName}Res." + svcrsbeplaceholder_text + "\n\t\t"  + placeholder + parts[1] 
                # modified_content = parts[0] + f"SET OutBody." + lines_to_add.split('.', 1)[-1]  + " = {BEName}Res." + lines_to_add.split('.', 1)[-1] + "\n\t\t"  + placeholder + parts[1] 
        else:
            if "LstItem" in lines_to_add:
                return
            else:
                modified_content = parts[0] + "SET ref." + lines_to_add.split(':', 1)[-1].replace('*', '').replace(':', '')  + " = in" + lines_to_add + "\n\t\t"  + placeholder + parts[1] 

        with open(output_file_path, "w") as output_file:
            output_file.write(modified_content)
        # print(f"Lines '{lines_to_add}' added after '{placeholder}' and placeholder removed. Output written to '{output_file_path}'.")
    else:
        print(f"Placeholder '{placeholder}' not found in the input file.")

def add_rows_after_placeholder(docx_file_path, placeholder, new_rows, cell1):
    doc = Document(docx_file_path)
    table_found = False

    # Iterate through all tables in the document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Check if the placeholder is found in any cell of the table
                if placeholder in cell.text:
                    table_found = True
                    # Get the index of the current row
                    current_row_index = row._element.getparent().index(row._element)
                    # Remove the placeholder text
                    cell.text = cell.text.replace(placeholder, "")
                    # Add new rows after the table row containing the placeholder
                    for _ in range(new_rows):
                        new_row_cells = table.add_row().cells
                        # Customize the content of the new rows if needed
                        # For example, you can set content for each cell in the new row:
                        new_row_cells[0].text = str(cell1)
                        new_row_cells[1].text = str(cell1).rsplit('/', 1)[-1] + "_Type"
                        if str(cell1)[-7:] == 'LstItem':
                            new_row_cells[2].text = "1:N"
                            new_row_cells[3].text = " "
                            new_row_cells[4].text = "R"
                        else:
                            new_row_cells[2].text = "0:1"
                            new_row_cells[3].text = " "
                            new_row_cells[4].text = "O"
                        new_row_cells[5].text = str(cell1).rsplit('/', 1)[-1] + placeholder

                    # Break after adding rows to avoid adding rows multiple times if the placeholder appears in multiple cells
                    break


    if not table_found:
        print("Placeholder not found in any table.")
    else:
        doc.save(docx_file_path)

def add_lines_after_placeholder(input_file_path, output_file_path, placeholder, lines_to_add, msgtype, value):
    global previousitem
    with open(input_file_path, "r") as input_file:
        file_content = input_file.read()

        if placeholder in file_content:
            parts = file_content.split(placeholder)
            if  msgtype == 'rs':
                if value == None and ("Lst" not in lines_to_add):
                    return
                elif str(lines_to_add)[-3:] == 'Lst' :
                    modified_content = parts[0] + "\n\t\t"  + f"DECLARE J INTEGER CARDINALITY({BEName}Res." + svcrsbeplaceholder_text  + "[]);\n\t\t" + "DECLARE I INTEGER 1;\n\n\t\t" + "CREATE FIELD OutBody."+ lines_to_add.rsplit('.', 1)[-1]  + ";\n\t\t" + "DECLARE outLst REFERENCE TO OutBody."+ lines_to_add.rsplit('.')[-1]+ ";\n\n\t\t"+"WHILE I <= J DO\n\t\t\t" + placeholder + parts[1]  
                elif str(lines_to_add)[-4:] == 'Item' :
                    if lines_to_add.rsplit('.', 1)[-1] in rsprocessed_elements:
                        return
                    # Mark the element as processed
                    rsprocessed_elements.add(lines_to_add.rsplit('.', 1)[-1])
                    modified_content = parts[0] + "\n\t\t\t" f"DECLARE inLst REFERENCE TO {BEName}Res."+ svcrsbeplaceholder_text + "[I]\n\t\t\t" + "CREATE LASTCHILD OF outLst NAME '" +lines_to_add.rsplit('.', 1)[-1] + "';\n\t\t\t"  + "DECLARE outLstItem REFERENCE TO '" + lines_to_add.rsplit('.', 1)[-1]+ "[<];\n\t\t\t" + placeholder + parts[1]  
                elif "LstItem" in lines_to_add:
                    if lines_to_add.rsplit('.', 1)[-1] in irsprocessed_elements:
                        return
                    # Mark the element as processed
                    irsprocessed_elements.add(lines_to_add.rsplit('.', 1)[-1])
                    previousitem = lines_to_add
                    modified_content = parts[0] + "SET outLstItem." + lines_to_add.rsplit('.', 1)[-1]  + " = inLst." + svcrsbeplaceholder_text + "\n\t\t\t"  + placeholder + parts[1] 
                elif "LstItem" in previousitem:
                    modified_content = parts[0] + "SET I = I+1;\n\t\tEND WHILE;\n\n\t\t" + f"SET OutBody." + lines_to_add.split('.', 1)[-1]  + " = {BEName}Res." + svcrsbeplaceholder_text + "\n\t\t" + placeholder + parts[1]
                    previousitem = 'globalizeditempreviously'
                else:
                    modified_content = parts[0] + f"SET OutBody." + lines_to_add.split('.', 1)[-1]  + " = {BEName}Res." + svcrsbeplaceholder_text + "\n\t\t"  + placeholder + parts[1] 
                    # modified_content = parts[0] + f"SET OutBody." + lines_to_add.split('.', 1)[-1]  + " = {BEName}Res." + lines_to_add.split('.', 1)[-1] + "\n\t\t"  + placeholder + parts[1] 
            else:
                if "LstItem" in lines_to_add:
                    return
                else:
                    modified_content = parts[0] + "SET ref." + lines_to_add.split(':', 1)[-1].replace('*', '').replace(':', '')  + " = in" + lines_to_add + "\n\t\t"  + placeholder + parts[1] 

            with open(output_file_path, "w") as output_file:
                output_file.write(modified_content)
            # print(f"Lines '{lines_to_add}' added after '{placeholder}' and placeholder removed. Output written to '{output_file_path}'.")
        else:
            print(f"Placeholder '{placeholder}' not found in the input file.")

def generate_complex_type(schema, field_name, xml_element):
    # Check if the element has already been processed, if yes, return
    if field_name in processed_elements:
        return
    # Mark the element as processed
    processed_elements.add(field_name)
    complex_type = ET.SubElement(schema, "xsd:complexType", name=f"{field_name}_Type")
    sequence = ET.SubElement(complex_type, "xsd:sequence")
    # Add a comment between elements (example comment text)
    # if ID == 1:
    #     comment = ET.Comment("#####################Complex Types####################")
    #     schema.append(comment)
    for child_element in xml_element:
        # if child_element.tag in iprocessed_elements and "Body_Type" not in child_element:
        #     return
        if len(child_element) > 0:
            # Handle child elements with LstItem types
            if child_element.tag[-7:] == 'LstItem':
                ET.SubElement(sequence, "xsd:element", name=child_element.tag, type=f"alr:{child_element.tag}_Type", maxOccurs="unbounded")
            # Handle child elements with complex types
            else: 
                ET.SubElement(sequence, "xsd:element", name=child_element.tag, type=f"alr:{child_element.tag}_Type", minoccurs="0")
            generate_complex_type(schema, child_element.tag, child_element)
            iprocessed_elements.add(child_element.tag)
        else:
            # Handle leaf nodes as simple types (assuming all leaf nodes are strings)
            if child_element.tag[-2:] == 'Dt':
                ET.SubElement(sequence, "xsd:element", name=child_element.tag, type="arb:Date_Type",  minOccurs="0")
            elif child_element.tag[-3:] == 'Cur':
                ET.SubElement(sequence, "xsd:element", name=child_element.tag, type="arb:Currency_Type",  minOccurs="0")
            elif child_element.tag[-4:] == 'Acct':
                ET.SubElement(sequence, "xsd:element", name=child_element.tag, type="arb:AcctNum_Type",  minOccurs="0")   
            elif 'Lang' in child_element.tag:
                ET.SubElement(sequence, "xsd:element", name=child_element.tag, type="arb:Lang_Type",  minOccurs="0")
            elif 'Name' in child_element.tag:
                ET.SubElement(sequence, "xsd:element", name=child_element.tag, type="arb:Name_Type",  minOccurs="0")
            elif child_element.tag[-3:] == 'Amt':
                ET.SubElement(sequence, "xsd:element", name=child_element.tag, type="arb:Amt_Type",  minOccurs="0")     
            elif 'Address' in child_element.tag:
                ET.SubElement(sequence, "xsd:element", name=child_element.tag, type="arb:Address_Type",  minOccurs="0")      
            else:
                ET.SubElement(sequence, "xsd:element", name=child_element.tag, type="arb:String_Type",  minOccurs="0")

def generate_xsd_element(parent, xml_element):
    xsd_element = ET.SubElement(parent, "xsd:element", name=xml_element.tag)
    xsd_element.set("type", f"alr:{xml_element.tag}_Type")


    generate_complex_type(parent, xml_element.tag, xml_element)
    processed_elements.clear()

def finetune_xml(parent=None):
    capitalize_xml_tags(parent)
    remove_invalid_suffix_t(parent)
    remove_invalid_suffix(parent)

def json_to_xml(json_data, parent=None):
    if isinstance(json_data, list):
        parent_tag = parent.tag + 'LstItem' if parent is not None else 'LstItem'
        for item in json_data:
            json_to_xml(item, parent=ET.SubElement(parent, parent_tag))
    elif isinstance(json_data, dict):
        for tag_name, value in json_data.items():
                element = ET.SubElement(parent, tag_name + 'Lst')
                json_to_xml(value, element)

    else:
        parent.text = str(json_data)

def doneedfuldeclartaions():
    global svcesqltemp_path
    os.makedirs(output_folder, exist_ok=True)
    os.makedirs(os.path.join(output_folder, sample_folder), exist_ok=True)
    if functionality >= 2:
        os.makedirs(os.path.join(output_folder, xsd_folder), exist_ok=True)
        os.makedirs(os.path.join(output_folder, wsdl_folder), exist_ok=True)
    if functionality >= 3:
        os.makedirs(os.path.join(output_folder, sci_folder), exist_ok=True)
        #define Document and replicate it
        wdoc_template = Document("NeededFiles/TemplateDocx.docx")
        wdoc_template.save(f"{output_folder}/{sci_folder}/{SvcID}-{SubSvcID}-SCI {OperationName}.docx")
    if functionality >= 4:
        os.makedirs(os.path.join(output_folder, app_folder), exist_ok=True)
        os.makedirs(os.path.join(output_folder, app_folder, svc_folder), exist_ok=True)
        os.makedirs(os.path.join(output_folder, app_folder, core_folder), exist_ok=True)
        os.makedirs(os.path.join(output_folder, app_folder, wsdl_folder), exist_ok=True)
        os.makedirs(os.path.join(output_folder, app_folder, xsd_folder), exist_ok=True)
        # Open the template file in read mode
        with open("NeededFiles/App/Template_Svc.esql", "r") as svcesqltemp_file:
            # Read the content of the template file
            svcesqltemp_template = svcesqltemp_file.read()
        svcesqltemp_path = f"{output_folder}/{app_folder}/{svc_folder}/{OperationName}_Svc.esql"
        # Open the output file in write mode and write the template content
        with open(svcesqltemp_path, "w") as svcesqltemp_path_out:
            svcesqltemp_path_out.write(svcesqltemp_template)

    if message_entry == 'paste':
        # Retrieve the request and response messages from form data
        request_message = request.form['Request']
        response_message = request.form['Response']
        if BECall in ('1','3'):
            # Save the pasted messages to files
            with open("NeededFiles/DesiredReq.xml", 'w') as req_file:
                req_file.write(request_message)
            with open("NeededFiles/DesiredRes.xml", 'w') as res_file:
                res_file.write(response_message)
        elif BECall == '2':
        # Save the pasted messages to files
            with open("NeededFiles/DesiredReq.json", 'w') as req_file:
                req_file.write(request_message)
            with open("NeededFiles/DesiredRes.json", 'w') as res_file:
                res_file.write(response_message)

def dosamples(): 
    global psamplereq
    global printsamplereq
    global printsampleres
    printsamplereq = format_xml(ET.tostring(xml_t_rq, encoding="utf-8").decode())
    with open("NeededFiles/TemplateSRQ.xml", "r") as samplerq:
            psamplereq = samplerq.read().format(SvcID=SvcID, SubSvcID=SubSvcID, OperationName=OperationName,
                                                CategoryName=CategoryName, ReqBody=printsamplereq)
    with open(f"{output_folder}/{sample_folder}/{OperationName}SampleReq.xml", "w") as file:
        # Write the content to the file using the write() method
        file.write(psamplereq)
    # sys.exit()

    #print sample response 
    global psampleres
    printsampleres = format_xml(ET.tostring(xml_t_rs, encoding="utf-8").decode())
    with open("NeededFiles/TemplateSRs.xml", "r") as samplers:
            psampleres = samplers.read().format(OperationName=OperationName,CategoryName=CategoryName, ResBody=printsampleres)
    # printsampleres.write(f"{output_folder}/{sample_folder}/{OperationName}SampleRes.xml")
    with open(f"{output_folder}/{sample_folder}/{OperationName}SampleRes.xml", "w") as file:
        # Write the content to the file using the write() method
        file.write(psampleres)

def dorequest():
    global xml_t_rq
    if BECall == '1':
        # Load the input XML file
        input_xml_path = 'NeededFiles/DesiredReq.xml'  # Replace with the actual path to your input XML file
        xml_t_rq = ET.parse(input_xml_path).getroot()
    elif BECall == '2':
        # Read JSON REQ
        with open('NeededFiles/DesiredReq.json', 'r') as jsonRq_file:
            json_Rq = json.load(jsonRq_file)
            #define request root
            xml_t_rq = ET.Element("Body")
            # Convert JSON to XML with a root element
            json_to_xml(json_Rq, xml_t_rq)
    elif BECall == '3':
        # Load the input XML file
        input_xml_path = 'NeededFiles/DesiredReq.xml'  # Replace with the actual path to your input XML file
        xml_t_rq = ET.parse(input_xml_path).getroot()
        for elem in xml_t_rq.iter():
            elem.tag = elem.tag.split('}')[-1]  # Remove namespaces
        xml_t_rq = xml_t_rq.find('.//Body/*')

            
    finetune_xml(xml_t_rq)
    #function 3: SCI generation
    if functionality >= 3:
    # Add records to SCI
        rqxml_sci_fields_iterator = islice(list_xml_fields_with_path(xml_t_rq), 1, None)
        for path, value in rqxml_sci_fields_iterator:
            add_rows_after_placeholder(f"{output_folder}/{sci_folder}/{SvcID}-{SubSvcID}-SCI {OperationName}.docx",
                                    Rqplaceholder_text, 1, path)
            # print(item)

    #function 4: SVC generation
    if functionality >= 4:
        # Add records to Svc Req Body
        rqxml_fields_iterator = islice(list_xml_fields_with_path_svc(xml_t_rq), 1, None)
        for path, value in rqxml_fields_iterator:
            add_lines_after_placeholder(svcesqltemp_path, svcesqltemp_path,
                                    svcrqplaceholder_text, path, 'rq', value)

def doresponse():
    global xml_t_rs
    if BECall == '1':
        # Load the input XML file
        inputr_xml_path = 'NeededFiles/DesiredRes.xml'  # Replace with the actual path to your input XML file
        xml_t_rs = ET.parse(inputr_xml_path).getroot()
    elif BECall == '2':
        with open('NeededFiles/DesiredRes.json', 'r') as jsonRs_file:
            json_Rs = json.load(jsonRs_file)
            #define request root
            xml_t_rs = ET.Element("Body")
            # Convert JSON to XML with a root element
            json_to_xml(json_Rs, xml_t_rs)
    elif BECall == '3':
        # Load the input XML file
        inputr_xml_path = 'NeededFiles/DesiredRes.xml'  # Replace with the actual path to your input XML file
        xml_t_rs = ET.parse(inputr_xml_path).getroot()
        for elem in xml_t_rs.iter():
            elem.tag = elem.tag.split('}')[-1]  # Remove namespaces
        xml_t_rs = xml_t_rs.find('.//Body/*')

    finetune_xml(xml_t_rs)


    #function 3: SCI generation
    if functionality >= 3:
        # Add records to SCI
        for path, value in list_xml_fields_with_path(xml_t_rs):
            add_rows_after_placeholder(f"{output_folder}/{sci_folder}/{SvcID}-{SubSvcID}-SCI {OperationName}.docx",
                                    placeholder_text, 1, path)
            # print(item)

    #function 4: SVC generation
    if functionality >= 4:
        # Add records to Svc Res Body service response
        rsxml_fields_iterator = islice(list_xml_fields_with_path_svcr(xml_t_rs), 1, None)
        for path, value in rsxml_fields_iterator:
            add_lines_after_placeholder(svcesqltemp_path, svcesqltemp_path,
                                    svcrsplaceholder_text, path, 'rs', value)

#endregion---------------------------functions-----------------------------------


app = Flask(__name__, static_url_path='/static')
    
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/submit', methods=['POST'])
def process_form():
    getuserinput()
    doneedfuldeclartaions()
    dorequest()
    doresponse() 
    #function 1: print sample request
    if functionality >= 1:
        dosamples()
        result =  "1:   samples printed succesfully<br>"
    #function 2: WSDL and XSD
    if functionality >= 2:
        dowsdlandxsd()
        result = result + "2:   XSD and WSDL Generated succesfully<br>"
        #function 3: SCI generation
    if functionality >= 3:  
        doworddoc()
        result = result + "3:   SCI Generated succesfully<br>"
        #function 4: SVC generation
    if functionality >= 4:
        dosvcandcoreapp()
        result = result + "4:   APP files Generated succesfully<br>"
    return render_template('result.html', result=result)    
   
@app.route('/samples')
def samples_page():
    return render_template('samples.html',samplerq = psamplereq,samplers = psampleres) 

@app.route('/doutput')
def sci_page():
    zip_folder(f"{OperationName}",f"{OperationName}.zip")
    return send_file(os.path.join(os.getcwd(), f"{OperationName}.zip"), as_attachment=True)

if __name__ == '__main__':
    app.run(host="0.0.0.0", port=5000)


